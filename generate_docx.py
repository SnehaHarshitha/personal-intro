import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_report():
    base_dir = os.path.dirname(os.path.abspath(__file__))
    visuals_dir = os.path.join(base_dir, 'reports/visuals')
    output_path = os.path.join(base_dir, 'reports/technical_report.docx')

    doc = docx.Document()

    # Title
    title = doc.add_heading('Technical Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Project Overview
    doc.add_heading('1. Project Overview', level=1)
    doc.add_paragraph(
        "This report detailing the end-to-end analysis of customer attrition drivers "
        "for a subscription-based service. The objective is to identify key risk factors "
        "and provide actionable business recommendations to reduce churn."
    )

    # Methodology
    doc.add_heading('2. Methodology', level=1)
    doc.add_paragraph(
        "The analysis followed a structured data science workflow including data cleaning, "
        "exploratory data analysis (EDA), and advanced statistical modeling using "
        "T-tests, Chi-Square analysis, and Logistic Regression."
    )

    # Data Insights
    doc.add_heading('3. Key Analytical Insights', level=1)
    
    # Insert Churn Distribution
    doc.add_heading('3.1 Churn Distribution', level=2)
    doc.add_paragraph("The overall churn distribution identifies the baseline attrition rate across the customer base.")
    img_path1 = os.path.join(visuals_dir, '1_churn_distribution.png')
    if os.path.exists(img_path1):
        doc.add_picture(img_path1, width=Inches(4.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Insert Tenure Impact
    doc.add_heading('3.2 Tenure and Attrition Risk', level=2)
    doc.add_paragraph("Analysis confirms that new customers within their first 90 days are at the highest risk.")
    img_path2 = os.path.join(visuals_dir, '2_tenure_vs_churn.png')
    if os.path.exists(img_path2):
        doc.add_picture(img_path2, width=Inches(4.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Statistical Results
    doc.add_heading('4. Statistical Findings', level=1)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Technique'
    hdr_cells[1].text = 'Metric'
    hdr_cells[2].text = 'Significance'
    
    row_cells = table.add_row().cells
    row_cells[0].text = 'Chi-Square'
    row_cells[1].text = 'p < 0.0001'
    row_cells[2].text = 'Highly Significant'

    # Recommendations
    doc.add_heading('5. Strategic Recommendations', level=1)
    doc.add_paragraph("1. Incentivize Long-term Contracts: Transition month-to-month users to annual plans.", style='List Bullet')
    doc.add_paragraph("2. Onboarding Engagement: Intensive proactive outreach for new users.", style='List Bullet')
    doc.add_paragraph("3. Revenue Tier Monitoring: Adjust high-billing service value to prevent attrition.", style='List Bullet')

    doc.save(output_path)
    print(f"Technical Report saved to {output_path}")

if __name__ == "__main__":
    create_report()
