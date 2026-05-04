import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_formal_document():
    doc = Document()
    
    # Title
    title = doc.add_heading('Week 12 Capstone Project: Progress Summary', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Introduction
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        "This document provides a formal summary of the activities, milestones, and technical achievements "
        "completed during the Week 12 Final Capstone project. The primary objective was to build an end-to-end "
        "machine learning pipeline focusing on Customer Churn Prediction, followed by deploying the solution "
        "as an interactive web service."
    )
    
    # Project Setup
    doc.add_heading('2. Project Initialization and Setup', level=1)
    doc.add_paragraph(
        "The project workspace was established following standard data science repository structures. "
        "Directories including 'src', 'data', 'reports', 'deployment', and 'presentation' were generated "
        "to ensure modularity and professional organization. A comprehensive README.md was created to guide "
        "reproducibility."
    )
    
    # Data Integration
    doc.add_heading('3. Data Integration and Preprocessing', level=1)
    doc.add_paragraph(
        "Real-world datasets were integrated into the project workflow. The primary dataset, 'customer_churn.csv', "
        "underwent rigorous preprocessing. Key steps included:\n"
        "• Removal of non-predictive identifiers (e.g., CustomerID).\n"
        "• One-hot encoding of categorical variables including 'Contract', 'PaymentMethod', and 'PaperlessBilling'.\n"
        "• Verification of target variable encoding (0 for retention, 1 for churn)."
    )
    
    # Model Development
    doc.add_heading('4. Model Training and Evaluation', level=1)
    doc.add_paragraph(
        "A Random Forest Classifier was selected as the predictive algorithm. The Jupyter Notebook "
        "('capstone_project.ipynb') was fully automated and executed to process the data, train the model, "
        "and output performance metrics. The classification results, including accuracy, precision, and recall, "
        "demonstrate the model's viability for business application."
    )
    
    # Deployment
    doc.add_heading('5. Model Deployment', level=1)
    doc.add_paragraph(
        "To operationalize the machine learning model, a robust deployment strategy was implemented:\n"
        "• Serialized the trained Random Forest model ('model.pkl').\n"
        "• Exported the precise feature schema ('columns.pkl') to guarantee compatibility between incoming API requests "
        "and the training environment.\n"
        "• Constructed a RESTful API using FastAPI ('app.py') that successfully exposes the '/predict' endpoint.\n"
        "• The server was successfully launched and hosted locally on http://127.0.0.1:8000, providing an interactive "
        "Swagger documentation interface for immediate testing."
    )
    
    # Conclusion
    doc.add_heading('6. Conclusion', level=1)
    doc.add_paragraph(
        "All phases of the data science lifecycle—from initial setup and data preprocessing to model training "
        "and real-time deployment—have been successfully executed. The system is currently live, effectively "
        "closing the loop from raw data to actionable business insights."
    )
    
    file_path = r"e:\dev\12_week\reports\Capstone_Project_Summary.docx"
    doc.save(file_path)
    print(f"Document successfully created at: {file_path}")

if __name__ == "__main__":
    create_formal_document()
